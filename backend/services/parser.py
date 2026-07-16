import io
import csv
from pypdf import PdfReader
from docx import Document as DocxDocument
from html.parser import HTMLParser
import requests

class HTMLTextExtractor(HTMLParser):
    """HTML Parser to strip tags and extract plain text."""
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text_parts = []
        self.in_ignored_tag = False
        self.ignored_tags = {"script", "style", "head", "title", "meta", "nav", "footer"}

    def handle_starttag(self, tag, attrs):
        if tag.lower() in self.ignored_tags:
            self.in_ignored_tag = True

    def handle_endtag(self, tag):
        if tag.lower() in self.ignored_tags:
            self.in_ignored_tag = False

    def handle_data(self, data):
        if not self.in_ignored_tag and data.strip():
            self.text_parts.append(data.strip())

    def get_text(self) -> str:
        return " ".join(self.text_parts)

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    extracted_text = []
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            extracted_text.append(text)
            
    return "\n".join(extracted_text)

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a Word Document (.docx)."""
    docx_file = io.BytesIO(file_bytes)
    doc = DocxDocument(docx_file)
    extracted_text = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            extracted_text.append(paragraph.text)
            
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                extracted_text.append(" | ".join(row_text))
                
    return "\n".join(extracted_text)

def parse_txt(file_bytes: bytes) -> str:
    """Extract text from a Plain Text file (.txt)."""
    return file_bytes.decode("utf-8", errors="ignore")

def parse_csv(file_bytes: bytes) -> str:
    """Extract text from a CSV file."""
    csv_text = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(csv_text))
    extracted_text = []
    
    for row in reader:
        # Join non-empty cell values
        cleaned_row = [cell.strip() for cell in row if cell.strip()]
        if cleaned_row:
            extracted_text.append(" ".join(cleaned_row))
            
    return "\n".join(extracted_text)

def parse_url(url: str) -> str:
    """Crawl a website URL and extract raw body text."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        parser = HTMLTextExtractor()
        parser.feed(response.text)
        text = parser.get_text()
        
        if not text.strip():
            raise Exception("No readable text content extracted from website.")
            
        return text
    except Exception as e:
        raise Exception(f"Failed to crawl URL: {str(e)}")

def extract_text_from_bytes(file_bytes: bytes, file_type: str) -> str:
    """Helper to route file bytes to correct parser based on file type extension."""
    ft = file_type.lower().strip(".")
    
    if ft == "pdf":
        return parse_pdf(file_bytes)
    elif ft in ["docx", "doc"]:
        return parse_docx(file_bytes)
    elif ft == "txt":
        return parse_txt(file_bytes)
    elif ft == "csv":
        return parse_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {file_type}")
